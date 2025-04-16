import os
from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# === CONFIG ===
PDF_PATH = "pdf/full_700_pages.pdf"
IMAGE_DIR = "images"
OUTPUT_DOCX = "output/cay_thuoc_an_giang.docx"
TESSERACT_CMD = r"tesseract\\tesseract.exe"  # chỉnh lại nếu đặt ở nơi khác
#TESSERACT_CMD = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
CHUNK_SIZE = 0  # = 0 để gộp tất cả vào một file duy nhất
LANG = "vie"

# === SETUP ===
pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD
os.makedirs(IMAGE_DIR, exist_ok=True)
os.makedirs("output", exist_ok=True)

# === CONVERT PDF TO IMAGES ===
print("🔄 Đang chuyển PDF thành ảnh (grayscale, dpi=400)...")
images = convert_from_path(
  PDF_PATH,
  dpi=400,
  grayscale=True,
  fmt="jpeg",
  output_folder=IMAGE_DIR,
  output_file="page",
)
print(f"✅ Đã tạo {len(images)} ảnh từ PDF")


# === OCR & EXPORT TO DOCX ===
def new_doc():
  d = Document()
  d.add_heading("Danh mục cây thuốc An Giang (OCR)", 0)
  add_page_number_footer(d)
  return d


def add_page_to_doc(doc, page_num, image_path, text):
  doc.add_page_break()
  doc.add_heading(f"Trang {page_num}", level=2)

  if not text.strip():
    doc.add_paragraph(
      "⚠️ Trang này có thể chỉ chứa hình vẽ hoặc ảnh, không có văn bản rõ ràng."
    )
  else:
    first_line = text.strip().split("\n")[0]
    if first_line.isupper() and len(first_line.split()) <= 8:
      doc.add_heading(first_line, level=1)
      text = text.replace(first_line, "", 1)
    doc.add_paragraph(text.strip())

  doc.add_picture(image_path, width=Inches(4.5))
  return doc


def save_doc(current_doc):
  current_doc.save(OUTPUT_DOCX)
  print(f"💾 Đã lưu file Word: {OUTPUT_DOCX}")


# Thêm đánh số trang vào footer


def add_page_number_footer(doc):
  section = doc.sections[0]
  footer = section.footer
  paragraph = footer.paragraphs[0]
  run = paragraph.add_run()

  fldChar1 = OxmlElement("w:fldChar")
  fldChar1.set(qn("w:fldCharType"), "begin")
  instrText = OxmlElement("w:instrText")
  instrText.text = "PAGE"
  fldChar2 = OxmlElement("w:fldChar")
  fldChar2.set(qn("w:fldCharType"), "end")

  run._r.append(fldChar1)
  run._r.append(instrText)
  run._r.append(fldChar2)


# === MAIN ===
print("🔍 Đang nhận diện văn bản bằng OCR...")
doc = new_doc()

# Ghi song song ra file .txt tổng hợp
text_output_path = OUTPUT_DOCX.replace(".docx", ".txt")
txt_file = open(text_output_path, "w", encoding="utf-8")

for i, image in enumerate(images):
  page_num = i + 1
  image_path = os.path.join(IMAGE_DIR, f"page_{page_num:04d}.jpg")
  if not os.path.exists(image_path):
    images[i].save(image_path)

  text = pytesseract.image_to_string(Image.open(image_path), lang=LANG)
  doc = add_page_to_doc(doc, page_num, image_path, text)

  txt_file.write(f"\n--- Trang {page_num} ---\n{text}\n")

# Lưu cuối cùng
save_doc(doc)
txt_file.close()
print(f"✅ Đã lưu file văn bản TXT song song: {text_output_path}")
print("✅ Hoàn tất xử lý OCR và xuất file Word!")
