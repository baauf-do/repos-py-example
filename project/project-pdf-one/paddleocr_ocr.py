import os
from pdf2image import convert_from_path
from paddleocr import PaddleOCR
from PIL import Image
from docx import Document
from docx.shared import Inches

# === CONFIG ===
PDF_PATH = "pdf/full_700_pages.pdf"
IMAGE_DIR = "images"
OUTPUT_DIR = "output"
DOCX_NAME_BASE = "cay_thuoc_part"
TXT_NAME = "cay_thuoc_ocr_full.txt"
PAGES_PER_DOCX = 100
LANG = "vi"

# === SETUP ===
os.makedirs(IMAGE_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
ocr = PaddleOCR(use_angle_cls=True, lang=LANG)

# === CONVERT PDF TO IMAGES ===
print("🔄 Đang chuyển PDF thành ảnh (dpi=400)...")
images = convert_from_path(
  PDF_PATH,
  dpi=400,
  fmt="jpeg",
  grayscale=True,
  output_folder=IMAGE_DIR,
  output_file="page",
)
print(f"✅ Đã tạo {len(images)} ảnh từ PDF")

# === DOCX EXPORT SETUP ===
all_text = []
doc_index = 1
page_index = 0
current_doc = Document()
current_doc.add_heading("Danh mục cây thuốc An Giang (OCR)", 0)


def add_heading_if_chapter(doc, line):
  # Nếu dòng viết HOA và ngắn, xem là tiêu đề chương
  if line.isupper() and 2 <= len(line.split()) <= 8:
    doc.add_heading(line.strip(), level=1)
    return True
  return False


# === XỬ LÝ TỪNG TRANG ===
for i, img in enumerate(images):
  page_num = i + 1
  image_path = os.path.join(IMAGE_DIR, f"page_{page_num:04d}.jpg")
  if not os.path.exists(image_path):
    img.save(image_path)

  print(f"📄 OCR Trang {page_num}...")
  result = ocr.ocr(image_path, cls=True)

  lines = []
  for line in result[0]:
    text = line[1][0].strip()
    if text:
      lines.append(text)

  page_text = "\n".join(lines)
  all_text.append(f"--- Trang {page_num} ---\n{page_text}\n")

  # --- Ghi vào file Word ---
  current_doc.add_page_break()
  current_doc.add_heading(f"Trang {page_num}", level=2)

  chapter_written = False
  if lines:
    first_line = lines[0]
    if add_heading_if_chapter(current_doc, first_line):
      chapter_written = True
      lines = lines[1:]  # Bỏ dòng đầu nếu đã ghi làm heading

    for para in lines:
      current_doc.add_paragraph(para)
  else:
    current_doc.add_paragraph(
      "⚠️ Trang này có thể chỉ chứa hình vẽ hoặc ảnh, không có văn bản rõ ràng."
    )

  # Chèn ảnh minh họa
  current_doc.add_picture(image_path, width=Inches(4.5))

  # --- Lưu sau mỗi 100 trang ---
  page_index += 1
  if page_index == PAGES_PER_DOCX or page_num == len(images):
    docx_path = os.path.join(OUTPUT_DIR, f"{DOCX_NAME_BASE}{doc_index}.docx")
    current_doc.save(docx_path)
    print(f"💾 Đã lưu: {docx_path}")
    doc_index += 1
    page_index = 0
    if page_num != len(images):
      current_doc = Document()
      current_doc.add_heading("Danh mục cây thuốc An Giang (OCR)", 0)

# === GHI FILE TXT TỔNG HỢP ===
txt_path = os.path.join(OUTPUT_DIR, TXT_NAME)
with open(txt_path, "w", encoding="utf-8") as f:
  f.writelines("\n".join(all_text))
print(f"📝 Đã xuất file văn bản tổng hợp: {txt_path}")

print("🎉 Hoàn tất toàn bộ quá trình OCR với PaddleOCR!")
print("📌 Mẹo: Vào Word > References > Table of Contents để tạo mục lục tự động!")
