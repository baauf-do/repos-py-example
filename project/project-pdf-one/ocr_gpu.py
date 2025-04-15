import os
from pdf2image import convert_from_path
from paddleocr import PaddleOCR
from docx import Document
from docx.shared import Inches
from io import BytesIO
import ctypes

try:
  ctypes.CDLL("zlibwapi.dll")
  print("✅ zlibwapi.dll đã được tìm thấy!")
except:
  print("❌ Thiếu zlibwapi.dll")
# === CONFIG ===
PDF_PATH = "pdf/full_700_pages.pdf"
OUTPUT_DIR = "output"
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, "cay_thuoc_an_giang.docx")
OUTPUT_TXT = os.path.join(OUTPUT_DIR, "cay_thuoc_an_giang.txt")
PAGES_PER_DOCX = 100  # Chia nhỏ file nếu cần
LANG = "vi"

# === SETUP ===
os.makedirs(OUTPUT_DIR, exist_ok=True)
ocr = PaddleOCR(use_angle_cls=True, lang=LANG, use_gpu=True)
print("🚀 PaddleOCR đã sẵn sàng (GPU)")

# === CHUYỂN PDF SANG ẢNH TRONG RAM ===
print("🔄 Đang chuyển PDF thành ảnh (grayscale, dpi=400)...")
images = convert_from_path(PDF_PATH, dpi=400, grayscale=True, fmt="jpeg")
print(f"✅ Đã tạo {len(images)} ảnh từ PDF (trong RAM)")


# === HÀM TẠO FILE DOCX MỚI ===
def new_doc():
  doc = Document()
  doc.add_heading("Danh mục cây thuốc An Giang (OCR)", 0)
  return doc


# === BẮT ĐẦU XỬ LÝ ===
doc = new_doc()
txt_lines = []
doc_index = 1
page_index = 0

for i, img in enumerate(images):
  page_num = i + 1
  print(f"🔍 OCR Trang {page_num}...")

  # OCR trực tiếp từ ảnh trong RAM (không lưu ra file)
  buffered = BytesIO()
  img.save(buffered, format="JPEG")
  result = ocr.ocr(buffered.getvalue(), cls=True)

  # Xử lý văn bản
  lines = [line[1][0].strip() for line in result[0] if line[1][0].strip()]
  page_text = "\n".join(lines)
  txt_lines.append(f"\n--- Trang {page_num} ---\n{page_text}\n")

  # Ghi vào DOCX
  doc.add_page_break()
  doc.add_heading(f"Trang {page_num}", level=2)

  if lines:
    first_line = lines[0]
    if first_line.isupper() and 2 <= len(first_line.split()) <= 8:
      doc.add_heading(first_line, level=1)
      lines = lines[1:]

    for para in lines:
      doc.add_paragraph(para)
  else:
    doc.add_paragraph(
      "⚠️ Trang này có thể chỉ chứa hình vẽ hoặc ảnh, không có văn bản rõ ràng."
    )

  # Lưu DOCX mỗi 100 trang hoặc cuối
  page_index += 1
  is_last_page = page_num == len(images)
  if page_index == PAGES_PER_DOCX or is_last_page:
    name = (
      f"cay_thuoc_an_giang_part{doc_index}.docx"
      if not is_last_page
      else "cay_thuoc_an_giang.docx"
    )
    path = os.path.join(OUTPUT_DIR, name)
    doc.save(path)
    print(f"💾 Đã lưu: {path}")
    doc_index += 1
    page_index = 0
    if not is_last_page:
      doc = new_doc()

# Ghi file TXT tổng hợp
with open(OUTPUT_TXT, "w", encoding="utf-8") as f:
  f.writelines(txt_lines)
print(f"📝 Đã xuất file văn bản tổng hợp: {OUTPUT_TXT}")

print("🎉 Hoàn tất toàn bộ quá trình OCR với PaddleOCR (GPU)!")
print(
  "📌 Gợi ý: Mở file Word và vào References > Table of Contents để tạo mục lục tự động."
)
